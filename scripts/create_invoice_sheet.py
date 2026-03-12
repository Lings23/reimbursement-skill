#!/usr/bin/env python3
"""
创建发票打印单：
- 第一页：PDF 转图片，尽可能占满页面
- 第二页：用餐时间和加班事由（四号字，居中）
"""

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Cm, Pt, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
import sys
from pathlib import Path


def create_invoice_sheet(input_dir, output_path, date='', reason=''):
    """创建发票打印单"""
    
    input_path = Path(input_dir)
    
    if not input_path.exists():
        print(f"❌ 目录不存在：{input_dir}")
        return
    
    pdf_files = list(input_path.glob('*.pdf'))
    
    if not pdf_files:
        print(f"⚠️ 目录中没有 PDF 文件：{input_dir}")
        return
    
    print(f"找到 {len(pdf_files)} 个 PDF 文件")
    
    # 创建文档
    doc = Document()
    
    # ========== 第一页：发票图片 ==========
    # 设置页面为纵向 A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)    # A4 宽度
    section.page_height = Inches(11.69)  # A4 高度
    
    # 设置页边距：上 4cm (1.57 英寸)，其他 2.54cm (1 英寸)
    section.top_margin = Inches(1.57)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # 可用尺寸（英寸）
    available_width = 6.27   # 8.27 - 1.0 - 1.0
    available_height = 9.12  # 11.69 - 1.57 - 1.0
    
    print(f"可用页面尺寸：{available_width:.2f}英寸 x {available_height:.2f}英寸")
    
    # 转换所有 PDF 为图片
    all_images = []
    for pdf_path in pdf_files:
        print(f"\n转换 PDF: {pdf_path.name}")
        try:
            images = convert_from_path(str(pdf_path), dpi=150)
            for i, img in enumerate(images):
                img_path = f"/tmp/invoice_page_{len(all_images)}.png"
                img.save(img_path, 'PNG')
                all_images.append(img_path)
                print(f"  - 保存页面：{img_path}")
        except Exception as e:
            print(f"  ⚠️ 转换失败：{e}")
    
    print(f"\n共 {len(all_images)} 张图片")
    
    if len(all_images) == 0:
        print("⚠️ 没有图片可插入")
        return
    
    # 计算布局
    num_images = len(all_images)
    if num_images == 1:
        cols = 1
        rows = 1
    elif num_images == 2:
        cols = 1
        rows = 2
    else:
        cols = 2
        rows = (num_images + 1) // 2
    
    spacing = 0.2  # 图片间距（英寸）
    
    # 计算每张图片的尺寸（英寸）
    img_width = (available_width - spacing * (cols - 1)) / cols
    img_height = (available_height - spacing * (rows - 1)) / rows
    
    print(f"布局：{rows}行 x {cols}列")
    print(f"每张图片：{img_width:.2f}英寸 x {img_height:.2f}英寸")
    
    # 插入图片（不添加标题）
    for idx, img_path in enumerate(all_images):
        row = idx // cols
        col = idx % cols
        
        # 获取图片尺寸
        with Image.open(img_path) as img:
            img_width_px, img_height_px = img.size
            img_ratio = img_width_px / img_height_px
            
            # 计算适应尺寸
            target_ratio = img_width / img_height
            
            if img_ratio > target_ratio:
                final_width = img_width
                final_height = img_width / img_ratio
            else:
                final_height = img_height
                final_width = img_height * img_ratio
        
        # 添加段落并插入图片
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(final_width), height=Inches(final_height))
        print(f"插入图片 {idx+1}: {final_width:.2f}英寸 x {final_height:.2f}英寸")
    
    # ========== 第二页：用餐信息 ==========
    doc.add_page_break()
    
    # 添加用餐信息（页面中间位置）
    for _ in range(10):  # 添加空行推到中间
        doc.add_paragraph()
    
    # 用餐时间
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"用餐时间：{date}")
    run_date.font.name = '宋体'
    run_date.font.size = Pt(14)  # 四号
    
    # 空行
    doc.add_paragraph()
    
    # 加班事由
    p_reason = doc.add_paragraph()
    p_reason.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_reason = p_reason.add_run(f"加班事由：{reason}")
    run_reason.font.name = '宋体'
    run_reason.font.size = Pt(14)  # 四号
    
    # 保存
    doc.save(output_path)
    print(f"\n✓ 文档已保存：{output_path}")
    return output_path


def main():
    import argparse
    parser = argparse.ArgumentParser(description='创建发票打印单')
    parser.add_argument('--input-dir', '-i', required=True, help='PDF 文件目录')
    parser.add_argument('--output', '-o', required=True, help='输出 docx 文件')
    parser.add_argument('--date', '-d', default='', help='用餐时间')
    parser.add_argument('--reason', '-r', default='', help='加班事由')
    
    args = parser.parse_args()
    
    create_invoice_sheet(args.input_dir, args.output, args.date, args.reason)


if __name__ == '__main__':
    main()
