#!/usr/bin/env python3
"""
使用模板填写加班餐费报销审批表
"""

import json
import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_cell_format(cell, font_name='宋体', font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格格式：宋体四号居中"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)


def fill_reimbursement(input_path, output_path, template_path, data):
    """填写报销审批表"""
    
    # 加载模板配置
    with open(template_path, 'r', encoding='utf-8') as f:
        template = json.load(f)
    
    print(f"✓ 已加载模板：{template['template_name']} v{template['version']}")
    
    # 打开文档
    doc = Document(input_path)
    table = doc.tables[template['document_structure']['table_index']]
    fields = template['fields']
    
    print(f"\n正在填写：{input_path}")
    
    # 填写用餐时间
    if '用餐时间' in data and '用餐时间' in fields:
        row = fields['用餐时间']['row']
        col = fields['用餐时间']['col']
        table.cell(row, col).text = data['用餐时间']
        set_cell_format(table.cell(row, col))
        print(f"  ✓ 用餐时间：{data['用餐时间']}")
    
    # 填写用餐人数
    if '用餐人数' in data and '用餐人数' in fields:
        row = fields['用餐人数']['row']
        col = fields['用餐人数']['col']
        table.cell(row, col).text = str(data['用餐人数'])
        set_cell_format(table.cell(row, col))
        print(f"  ✓ 用餐人数：{data['用餐人数']}")
    
    # 填写用餐金额
    if '用餐金额' in data and '用餐金额' in fields:
        row = fields['用餐金额']['row']
        col = fields['用餐金额']['col']
        table.cell(row, col).text = str(data['用餐金额'])
        set_cell_format(table.cell(row, col))
        print(f"  ✓ 用餐金额：{data['用餐金额']} 元")
    
    # 填写加班事由
    if '加班事由' in data and '加班事由' in fields:
        row = fields['加班事由']['row']
        col = fields['加班事由']['col']
        table.cell(row, col).text = data['加班事由']
        set_cell_format(table.cell(row, col))
        print(f"  ✓ 加班事由：{data['加班事由']}")
    
    # 填写用餐人员名单
    if '用餐人员名单' in data and '用餐人员名单' in fields:
        names = data['用餐人员名单']
        name_config = fields['用餐人员名单']
        
        name_cols = name_config['name_columns']
        name_index = 0
        
        for col_config in name_cols:
            col = col_config['col']
            max_items = col_config['max_items']
            start_row = name_config['data_start_row']
            
            for i in range(max_items):
                if name_index < len(names):
                    row = start_row + i
                    table.cell(row, col).text = names[name_index]
                    set_cell_format(table.cell(row, col))
                    print(f"  ✓ 姓名 [{row},{col}]: {names[name_index]}")
                    name_index += 1
    
    # 保存
    doc.save(output_path)
    print(f"\n✓ 文档已保存：{output_path}")
    return output_path


def parse_names(names_str):
    """解析姓名列表（支持多种分隔符）"""
    if isinstance(names_str, list):
        return names_str
    
    # 替换常见分隔符为逗号（包括中文逗号、顿号、空格、分号）
    names_str = names_str.replace('，', ',').replace('、', ',').replace(' ', ',').replace(';', ',')
    names = [n.strip() for n in names_str.split(',') if n.strip()]
    return names


def main():
    parser = argparse.ArgumentParser(description='填写加班餐费报销审批表')
    parser.add_argument('--input', '-i', required=True, help='输入 docx 文件')
    parser.add_argument('--output', '-o', required=True, help='输出 docx 文件')
    parser.add_argument('--template', '-t', required=True, help='模板 JSON 文件')
    parser.add_argument('--names', '-n', required=True, help='用餐人员名单（逗号分隔）')
    parser.add_argument('--amount', '-a', type=int, required=True, help='用餐金额')
    parser.add_argument('--reason', '-r', default='', help='加班事由')
    parser.add_argument('--date', '-d', default='', help='用餐时间')
    
    args = parser.parse_args()
    
    # 准备数据
    names = parse_names(args.names)
    data = {
        '用餐人数': len(names),
        '用餐金额': args.amount,
        '用餐人员名单': names,
        '加班事由': args.reason,
        '用餐时间': args.date
    }
    
    # 填写文档
    fill_reimbursement(args.input, args.output, args.template, data)


if __name__ == '__main__':
    main()
