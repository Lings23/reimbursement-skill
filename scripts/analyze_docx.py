#!/usr/bin/env python3
"""
分析 docx 文档并自动生成 JSON 模板
"""

import json
from docx import Document
from pathlib import Path


def analyze_table(table):
    """分析表格结构"""
    structure = {
        "total_rows": len(table.rows),
        "total_cols": len(table.columns),
        "cells": []
    }
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                structure["cells"].append({
                    "row": row_idx,
                    "col": col_idx,
                    "text": cell_text,
                    "type": detect_field_type(cell_text)
                })
    
    return structure


def detect_field_type(text):
    """根据单元格文本检测字段类型"""
    type_keywords = {
        "number": ["人数", "数量", "序号"],
        "currency": ["金额", "费用", "元", "¥"],
        "date": ["日期", "时间", "年", "月"],
        "text": ["部门", "账号", "事由", "备注"],
        "name": ["姓名", "名单", "人员"]
    }
    
    for field_type, keywords in type_keywords.items():
        if any(kw in text for kw in keywords):
            return field_type
    
    return "unknown"


def generate_template(doc_path, output_path=None):
    """生成模板 JSON"""
    doc = Document(doc_path)
    
    # 分析第一个表格
    if len(doc.tables) == 0:
        print("❌ 文档中没有表格")
        return None
    
    table = doc.tables[0]
    table_structure = analyze_table(table)
    
    # 自动识别字段
    fields = {}
    name_list_config = None
    
    for cell in table_structure["cells"]:
        text = cell["text"]
        row = cell["row"]
        col = cell["col"]
        
        # 识别标签单元格及其对应的值单元格
        if "用餐人数" in text:
            fields["用餐人数"] = {
                "type": "number",
                "row": row,
                "col": col + 1,
                "label_cell": {"row": row, "col": col}
            }
        elif "用餐金额" in text:
            fields["用餐金额"] = {
                "type": "currency",
                "row": row,
                "col": col + 1,
                "label_cell": {"row": row, "col": col},
                "unit": "元"
            }
        elif "用餐时间" in text and "支出账号" not in text:
            fields["用餐时间"] = {
                "type": "date",
                "row": row,
                "col": col + 1,
                "label_cell": {"row": row, "col": col},
                "format": "YYYY年M月D日"
            }
        elif "支出账号" in text:
            fields["支出账号"] = {
                "type": "text",
                "row": row,
                "col": col + 1,
                "label_cell": {"row": row, "col": col}
            }
        elif "加班事由" in text:
            fields["加班事由"] = {
                "type": "text",
                "row": row,
                "col": col + 1,
                "label_cell": {"row": row, "col": col},
                "merge_cols": list(range(col + 1, table_structure["total_cols"]))
            }
        elif "用餐人员名单" in text and col == 0:
            # 名单区域
            name_list_config = {
                "type": "name_list",
                "header_row": row + 1,  # 序号/姓名标题行
                "data_start_row": row + 2,  # 数据开始行
                "data_end_row": table_structure["total_rows"] - 4,  # 审批行之前
                "name_columns": []
            }
        elif text == "姓名" and name_list_config:
            # 姓名列
            name_list_config["name_columns"].append({
                "col": col,
                "seq_col": col - 1,
                "max_items": 4
            })
        elif "部门" in text and col == 0:
            fields["部门"] = {
                "type": "text",
                "row": row,
                "col": col + 1,
                "default": ""
            }
    
    # 添加名单配置
    if name_list_config and name_list_config["name_columns"]:
        fields["用餐人员名单"] = name_list_config
    
    # 构建完整模板
    template = {
        "template_name": "加班餐费报销审批表",
        "version": "2.0",
        "created": "2026-03-11",
        "source_doc": str(doc_path),
        "description": "从文档自动生成的模板",
        
        "document_structure": {
            "table_index": 0,
            "total_rows": table_structure["total_rows"],
            "total_cols": table_structure["total_cols"]
        },
        
        "fields": fields,
        
        "validation_rules": {
            "用餐人数": {
                "min": 1,
                "max": 12,
                "must_match_name_count": True
            },
            "用餐金额": {
                "min": 0,
                "max": 10000,
                "currency": "CNY"
            }
        },
        
        "auto_fill": {
            "用餐人数": {
                "source": "name_list_count",
                "description": "自动计算姓名数量"
            }
        }
    }
    
    # 保存模板
    if output_path is None:
        output_path = str(doc_path).replace('.docx', '_template.json')
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(template, f, ensure_ascii=False, indent=4)
    
    print(f"✓ 模板已生成：{output_path}")
    print(f"\n识别到的字段：")
    for field_name, config in fields.items():
        print(f"  - {field_name}: 行{config['row']}, 列{config['col']}")
    
    return template


if __name__ == '__main__':
    import sys
    
    # 获取脚本所在目录
    SCRIPT_DIR = Path(__file__).parent
    SKILL_DIR = SCRIPT_DIR.parent
    
    # 默认分析模板文件（使用相对路径）
    doc_path = sys.argv[1] if len(sys.argv) > 1 else str(SKILL_DIR / 'templates/加班餐费报销审批表模板.docx')
    
    if not Path(doc_path).exists():
        # 尝试在当前目录查找
        local_template = Path('templates/加班餐费报销审批表模板.docx')
        if local_template.exists():
            doc_path = str(local_template)
            print(f"使用本地模板：{doc_path}")
        else:
            print(f"❌ 模板文件不存在：{doc_path}")
            print("用法：python analyze_docx.py <模板 docx 文件>")
            sys.exit(1)
    
    generate_template(doc_path)
